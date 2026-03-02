import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Modify default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Avances de la Solución Informática - Proyecto de Comercio Electrónico', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Materia: Comercio Electrónico')
    doc.add_paragraph('Actividad: Semana 5, 6 y 7 - Proyecto Avances')
    doc.add_paragraph('Descripción: Capturas de la solución informática y explicación de sus avances.')
    doc.add_paragraph('') # Empty space

    screenshots_dir = os.path.join(os.path.dirname(__file__), 'screenshots_tmp')
    
    advances = [
        {
            'file': '1_login.png',
            'title': '1. Pantalla de Inicio de Sesión (Login)',
            'desc': 'Se ha implementado la pantalla de inicio de sesión con validación de credenciales en conexión con el backend (API REST Node.js). El acceso está restringido según los roles del usuario (Administrador, Cajero, etc.).'
        },
        {
            'file': '2_login_filled.png',
            'title': '2. Ingreso de Credenciales',
            'desc': 'El formulario cuenta con validación reactiva de Angular, asegurando que los campos requeridos y el formato del correo sean correctos antes de enviar la petición.'
        },
        {
            'file': '3_dashboard.png',
            'title': '3. Panel de Control (Dashboard)',
            'desc': 'Una vez autenticado el acceso al sistema, el usuario visualiza el Dashboard principal donde se presenta un resumen de las métricas clave, además del menú lateral de navegación hacia los distintos módulos del sistema.'
        },
        {
            'file': '4_productos.png',
            'title': '4. Módulo de Productos (Inventario)',
            'desc': 'Se desarrolló el módulo de inventario que permite visualizar, crear y editar los productos existentes, integrándose en tiempo real con la base de datos MySQL.'
        },
        {
            'file': '5_ventas.png',
            'title': '5. Módulo de Ventas',
            'desc': 'La sección de ventas o pos permite el registro de las transacciones. Actualmente soporta cálculo de totales y almacenamiento de la orden para el respectivo control de caja.'
        },
        {
            'file': '6_usuarios.png',
            'title': '6. Gestión de Usuarios y Roles',
            'desc': 'Implementación del módulo de administración de usuarios que permite crear perfiles de acceso y definir roles en el aplicativo, garantizando la seguridad en las operaciones.'
        }
    ]

    for adv in advances:
        img_path = os.path.join(screenshots_dir, adv['file'])
        
        doc.add_heading(adv['title'], level=2)
        doc.add_paragraph(adv['desc'])
        
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f'[Error al insertar imagen: {str(e)}]')
        else:
            doc.add_paragraph('[Imagen no disponible]')
        
        doc.add_page_break()

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Avances_Proyecto_Comercio_v2.docx')
    doc.save(output_path)
    print(f'Document saved successfully to {output_path}')

if __name__ == '__main__':
    main()
